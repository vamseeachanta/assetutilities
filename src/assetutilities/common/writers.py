# ABOUTME: Data persistence utilities for JSON, YAML, CSV, Excel, and ASCII formats.
# ABOUTME: Extracted from common/data.py SaveData class.
from __future__ import annotations

import logging
from typing import Any, Optional

import numpy as np
import pandas as pd
import yaml


class SaveData:
    def saveDataJson(self, data: Any, fileName: str) -> None:
        import json

        with open(fileName + ".json", "w") as f:
            json.dump(data, f)

    def saveDataYaml(self, data: Any, fileName: str, default_flow_style: Any = False, sort_keys: bool = False) -> None:
        import yaml

        if default_flow_style is None:
            with open(fileName + ".yml", "w") as f:
                yaml.dump(data, f)

        elif default_flow_style == "NonAlias":
            with open(fileName + ".yml", "w") as f:
                yaml.dump(data, f, Dumper=noalias_dumper)  # type: ignore[name-defined]
        elif default_flow_style == "ruamel":
            from ruamel.yaml import ruamel

            with open(fileName + ".yml", "w") as f:
                ruamel.yaml.dump(data, f)
        elif default_flow_style == "round_trip_dump":
            with open(fileName + ".yml", "w") as f:
                ruamel.yaml.round_trip_dump(data, f)  # type: ignore[used-before-def]
        else:
            file_path = fileName + ".yml"
            with open(file_path, "w") as f:
                yaml.dump(
                    data, f, default_flow_style=default_flow_style, sort_keys=sort_keys
                )

    def saveDataFrame(self, df: pd.DataFrame, fileName: str) -> None:
        df.to_csv(fileName + ".csv")

    def save_ascii_file_from_array(self, array: list[Any], file_name: str, extension: str = "") -> None:
        with open(file_name + extension, "w") as f:
            for line_index in range(0, len(array)):
                if type(array[line_index]) is list:
                    from common.ETL_components import ETL_components

                    etl_components = ETL_components(cfg=None)
                    array[line_index] = (
                        etl_components.convert_text_list_to_combined_text(
                            array[line_index]
                        )
                    )
                f.write(array[line_index])
                if "\n" not in array[line_index]:
                    f.write("\n")

    def DataFrame_To_xlsx_xlsxwriter(df: Any, data: dict[str, Any]) -> None:
        """
        https://xlsxwriter.readthedocs.io/working_with_pandas.html
        """
        writer = pd.ExcelWriter(data["FileName"], engine="xlsxwriter")
        wb = writer.book

        try:
            wb[data["SheetName"]]
        except:
            wb.create_sheet(data["SheetName"])
            wb[data["SheetName"]]

        df.to_excel(writer, data["SheetName"])
        writer.save()

    """
    References:
    https://stackoverflow.com/questions/42370977/how-to-save-a-new-sheet-in-an-existing-excel-file-using-pandas/42371251
    https://stackoverflow.com/questions/36814050/openpyxl-get-sheet-by-name
    """

    def DataFrame_To_xlsx_openpyxl(self, df: pd.DataFrame, data: dict[str, Any]) -> None:
        import pandas as pd
        from openpyxl import load_workbook

        try:
            wb = load_workbook(data["FileName"])
            writer = pd.ExcelWriter(data["FileName"], engine="openpyxl")
            writer.wb = wb
        except:
            writer = pd.ExcelWriter(data["FileName"])

        try:
            wb[data["SheetName"]]
        except:
            wb.create_sheet(data["SheetName"])
            wb[data["SheetName"]]

        df.to_excel(writer, data["SheetName"])
        writer.save()

    def DataFrameArray_To_xlsx_openpyxl(self, dfArray: list[pd.DataFrame], data: dict[str, Any]) -> None:
        import pandas as pd

        print("Opening new workbook")
        writer = pd.ExcelWriter(data["FileName"], engine="openpyxl")

        for dfIndex in range(0, len(dfArray)):
            dfArray[dfIndex].to_excel(writer, sheet_name=data["SheetNames"][dfIndex])

            """
            References:
            https://stackoverflow.com/questions/24917201/applying-borders-to-a-cell-in-openpyxl
            """
            # property cell.border should be used instead of cell.style.border
            # if data['thin_border']:
            #     ws = wb.active
            #     thin_border = Border(
            #         left=Side(border_style=BORDER_THIN, color='00000000'),
            #         right=Side(border_style=BORDER_THIN, color='00000000'),
            #         top=Side(border_style=BORDER_THIN, color='00000000'),
            #         bottom=Side(border_style=BORDER_THIN, color='00000000')
            #     )
            #     rows = len(dfArray[dfIndex])+1
            #     columns = len(dfArray[dfIndex].columns)+1
            #     ws.cell(row=rows, column=columns).border = thin_border

        writer._save()

    def df_to_sheet_in_existing_workbook(self, cfg: dict[str, Any]) -> None:
        # cfg_example = {'template_file_name': file_name, 'sheetname': sheetname, 'saved_file_name': file_name, 'if_sheet_exists': 'replace', 'df': df, 'index': False}

        template_file_name = cfg["template_file_name"]
        cfg["saved_file_name"]
        sheetname = cfg["sheetname"]
        df = cfg["df"]
        if_sheet_exists = cfg.get("if_sheet_exists", "replace")
        index = cfg.get("index", True)

        if len(df) > 0:
            with pd.ExcelWriter(
                template_file_name,
                engine="openpyxl",
                if_sheet_exists=if_sheet_exists,
                mode="a",
            ) as writer:
                df.to_excel(writer, sheet_name=sheetname, index=index)
            logging.info(
                "Data write to sheet {sheetname} {Fore.GREEN}    SUCCESS... {Style.RESET_ALL}"
            )
        else:
            logging.info("{Fore.RED}No data to write to sheet{Style.RESET_ALL}")

    def df_to_table_as_image(self, df: pd.DataFrame, cfg: dict[str, Any]) -> None:
        (ax, fig) = self.render_df_table(
            df, header_columns=0, col_width=2.0, font_size=8
        )
        fig.savefig(cfg["file_name"])
        self.plt.close()

    def df_to_table_as_docx(self, df: pd.DataFrame, cfg: dict[str, Any]) -> None:
        from docx import Document

        document = Document()
        document.add_heading("Heading")
        table = document.add_table(
            rows=(df.shape[0] + 1), cols=df.shape[1]
        )  # First row are table headers!

        for j in range(df.shape[-1]):
            table.cell(0, j).text = df.columns[j]
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                table.cell(i + 1, j).text = str(df.values[i, j])

        document.save(cfg["file_name"] + ".docx")

    def render_df_table(
        self,
        data: pd.DataFrame,
        col_width: float = 3.0,
        row_height: float = 0.625,
        font_size: int = 14,
        header_color: str = "#40466e",
        row_colors: Optional[list[str]] = None,
        edge_color: str = "w",
        bbox: Optional[list[float]] = None,
        header_columns: int = 0,
        ax: Any = None,
        **kwargs: Any,
    ) -> Any:
        import matplotlib.pyplot as plt
        import six

        if bbox is None:
            bbox = [0, 0, 1, 1]
        if row_colors is None:
            row_colors = ["#f1f1f2", "w"]
        self.plt = plt

        if ax is None:
            size = (np.array(data.shape[::-1]) + np.array([0, 1])) * np.array(
                [col_width, row_height]
            )
            fig, ax = plt.subplots(figsize=size)
            ax.axis("off")

        mpl_table = ax.table(
            cellText=data.values, bbox=bbox, colLabels=data.columns, **kwargs
        )

        mpl_table.auto_set_font_size(False)
        mpl_table.set_fontsize(font_size)

        for k, cell in six.iteritems(mpl_table._cells):
            cell.set_edgecolor(edge_color)
            if k[0] == 0 or k[1] < header_columns:
                cell.set_text_props(weight="bold", color="w")
                cell.set_facecolor(header_color)
            else:
                cell.set_facecolor(row_colors[k[0] % len(row_colors)])
        return ax, fig

    def write_ascii_file_from_text(self, text: str, file_name: str) -> None:
        f = open(file_name, "wb")
        f.write(text.encode())
        f.close()

    def WriteOrcaFlexYMLFile(self, Files: list[Any], output_filename: str) -> None:
        print(f'Write file: "{output_filename}" .... ')
        with open(output_filename, "w") as f:
            for file in Files:
                if isinstance(file, str):
                    with open(file) as infile:
                        f.write(infile.read())
                else:
                    yaml.dump(file, f)
        print(f'Write file: "{output_filename}" .... COMPLETE')
